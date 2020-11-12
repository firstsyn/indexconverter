#!/usr/bin/env python3 -tt
#-*- coding: UTF-8 -*-
"""Convert Excel .csv file to Word .docx file.

Convert an index for a course (i.e., SANS) from Excel to Word.  The first
step is to save the main Excel index spreadsheet from the .xlsx file to
a UTF-8 comma-delimited .csv file; this step is done manually in Excel.
This script then converts the .csv file to a Word .docx file in a
layout that can be easily printed and used during the open book cert exam.

Parameters
----------
csvfile : str
  Name of index CSV file to convert to .docx.

Output
------
<csvfile>.docx
  Word document created from CSV index file; created in directory script is run from.

Requires
--------
python-docx
  https://python-docx.readthedocs.io/en/latest/
  https://buildmedia.readthedocs.org/media/pdf/python-docx/latest/python-docx.pdf

Format
------
csvfile
  Comma-separated data columns: TOPIC, BK#, PG#, COMMENTS
"""


# Future imports.
from __future__ import print_function

# Standard library imports.
import argparse
import csv
import os
import sys

# Third party imports.
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# Documentation (per PEP 8, 20, 257, PEP 263).
__author__ = "Frank Vasko"
__copyright__ = "Copyright 2020"
__date__ = "2020/11/11"
__license__ = "GPLv3"
__maintainer__ = "Frank Vasko"
__status__ = "Production"
__usage__ = "indexconverter.py [-h] csvfile"
__version__ = "1.0"


# If Python2, sets input to raw_input.
try:
    input = raw_input
except:
    pass


def add_page_numbers(paragraph):
    '''Add page numbers to a document.

    Parameters
    ----------
    paragraph : Document (Document().sections[].footer.paragraphs[])
      The section footer paragraph to add page numbers to.

    Reference
    ---------
    Code lifted from here (with edits):
    https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx
    '''

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = 'Page '
    run._r.append(t)

    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

    run = paragraph.add_run()
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = ' of '
    run._r.append(t)

    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)


def main():

    # Sort the CSV file; sorts all non-alpha characters to the top;
    # ignore case for alpha characters.  Known issue:  does not sort
    # {, |, }, ~ before alpha characters due to ASCII number.
    if sys.version_info.major == 2:
        csvin = csv.reader(open(args.csvfile))
        csvsorted = sorted(csvin, key=lambda row:(row[0].lower(), row[1], row[2], row[3].lower()))
    else:
        csvin = csv.reader(open(args.csvfile, encoding='UTF8'))
        csvsorted = sorted(csvin, key=lambda row:(row[0].casefold(), row[1], row[2], row[3].casefold()))

    # Set the .docx filename and check if it exists.
    docfile = os.path.splitext(args.csvfile)[0]+".docx"
    if os.path.isfile(docfile):
        parser.error(docfile+' already exists.')

    # Create the .docx document object.
    document = Document()

    # Set the document section settings, including page numbers
    section = document.sections[0]
    add_page_numbers(section.footer.paragraphs[0])
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

    # Set the document heading (Title) settings
    style = document.styles['Title']
    font = style.font
    font.size = Pt(32)
    font.bold = True
    #font.name = 'Times New Roman'  # Does not actually set the font name…
    rFont = style.element.rPr.rFonts  # …so, use XML setting to set the font name.
    rFont.set(qn('w:asciiTheme'), 'Times New Roman')
    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Inches(0)
    paragraph_format.first_line_indent = Inches(-0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set the document Body (Normal) settings.
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Inches(0.1)  # Hanging indent require a left indent…
    paragraph_format.first_line_indent = Inches(-0.1)  # …and then a first line negative indent.
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Read the rows from the CSV file and put in the document object.
    currltr = "`"  # Used to track the title heading.
    paragraph = document.add_heading("#", level=0)
    print("Processing:  #")
    for row in csvsorted:
        # Check if new letter; if so, print new title heading.
        if row[0][0].lower() > currltr:
            print("Processing:  "+row[0][0].upper()+row[0][0].lower())
            currltr = row[0][0].lower()
            #document.add_page_break()  # Use this if you want to start each "alpha" on just a new page…
            section = document.add_section(WD_SECTION.ODD_PAGE)  # …and comment out this line.
            paragraph = document.add_heading(row[0][0].upper()+row[0][0].lower(), level=0)
        paragraph = document.add_paragraph()
        run = paragraph.add_run(row[0])
        run.bold = True
        run.font.color.rgb = RGBColor(22, 103, 255)
        paragraph.add_run(" [b"+row[1]+"/p"+row[2]+"] ").italic = True
        if sys.version_info.major == 2:
            paragraph.add_run(row[3].decode('utf-8'))
        else:
            paragraph.add_run(row[3])

    # Save the document object to a Word .docx file in the current directory.
    document.save(docfile)
    print("Document created:  "+docfile)


if __name__ == "__main__":
    # Check the arguments passed to the program.
    parser = argparse.ArgumentParser()
    parser.add_argument('csvfile', help='name of index CSV file to convert to .docx')
    args = parser.parse_args()
    if not os.path.isfile(args.csvfile):
        parser.error('CSV file does not exist.')
    if os.path.getsize(args.csvfile) == 0:
        parser.error('CSV file is empty.')
    main()